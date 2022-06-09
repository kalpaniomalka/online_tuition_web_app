import docx
import os 

current_path = os.path.abspath(os.path.join(os.path.dirname(__file__)))
summarized_doc_path =  current_path+"\summarized_doc" 
fileName = "originalText.docx"
accuracy_check_file = "accuracy_check.txt"

# read system summarized doc
summarized_doc = docx.Document()
text = ""
doc = docx.Document(summarized_doc_path+'\\'+fileName)
AllLines = []
AllWords = []
for para in doc.paragraphs:
    AllWords.append((para.text).split(" "))

# read reference summary
with open(accuracy_check_file, 'r') as file:
    content = file.read()

no_of_overlapping_words = 0

for word in AllWords[0]:
    if word in content:
        no_of_overlapping_words += 1

print(no_of_overlapping_words)
total_words_in_system_summery = len(AllWords[0])
print(total_words_in_system_summery)

content_words = []
content_words.append(content.split(" "))
total_words_in_reference_summery = len(content_words[0])

# In terms of precision, what you are essentially measuring is, how much of the system summary 
# was in fact relevant or needed?
precision = no_of_overlapping_words / total_words_in_system_summery

# Recall in the context of ROUGE means how much of the reference summary is the system 
# summary recovering or capturing
recall = no_of_overlapping_words / total_words_in_reference_summery

print("Precision and Recall in the Context of ROUGE")
print("--------------------------------------------")
print("Recall: "+str(recall))
print("Text Summerization Accuracy(Precision): "+str(precision))
