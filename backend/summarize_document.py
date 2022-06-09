import spacy
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation
from heapq import nlargest
import os
import docx

# path declaration
current_path = os.path.abspath(os.path.join(os.path.dirname(__file__)))
document_path =  current_path+"\\documents" 
summarized_doc_path =  current_path+"\\summarized_doc"

summarized_doc = docx.Document()

# get list of stop words
stopwords = list(STOP_WORDS)
# get list of punctuations and add new line(\n) to the list 
punctuation = punctuation + '\n'

def summerize(fileName):
 
    fn = fileName.split(".")

    text = ""

    doc = docx.Document(document_path+'\\'+fileName)
    AllLines = []

    # Read content of the file and store in AllLines 
    for para in doc.paragraphs:
        AllLines.append(para.text)

    # Loop through AllLines and remove both the leading and the trailing characters 
    for line in AllLines:
        text = text + line.strip()

    # load spacy (spacy - data analysis tool in NLP)
    nlp = spacy.load('en_core_web_sm')
    doc = nlp(text)

    # apply tokenization
    tokens = [token.text for token in doc]
    print("Tokens: \n"+str(tokens))

    # text cleaning - remove stop words and punctuations
    # obtaining the frequency of words that are not considered as stop words (a, the, is, ..)
    word_frequency = {}
    for word in doc:
        if word.text.lower() not in stopwords:
            if word.text.lower() not in punctuation:
                if word.text not in word_frequency.keys():
                    word_frequency[word.text] = 1
                else: #already exist
                    word_frequency[word.text] += 1
    print("\nWords Frequencies: \n"+str(word_frequency))

    # get max frequency
    max_frequency = max(word_frequency.values())
    print("\nMax Frequency: "+str(max_frequency))

    # to get normalized word frequencies, divide word frequency by max frequency
    for word in word_frequency.keys():
        word_frequency[word] = word_frequency[word]/max_frequency
    print("\nWord Frequncies after devided by max frequency: \n"+str(word_frequency))

    # get sentence tokens
    sentence_tokens = [sent for sent in doc.sents]
    print("\nSentences tokens:\n"+str(sentence_tokens))

    # give scores to each sentences, by adding the word frequencies in each sentence
    sentence_score = {}
    for sent in sentence_tokens:
        for word in sent:
            if word.text.lower() in word_frequency.keys():
                if sent not in sentence_score.keys():
                    sentence_score[sent] = word_frequency[word.text.lower()]
                else:
                    sentence_score[sent] += word_frequency[word.text.lower()]
    print("\nSentences Scores: \n"+str(sentence_score))

    # calculate  40% of text with maximum score
    select_length = int(len(sentence_tokens)*0.4)
    print("\nSelect Length: \n"+str(select_length))

    # create summary
    summary = nlargest(select_length, sentence_score, key = sentence_score.get)

    print("\nSummary:\n"+str(summary))

    # get the final summary
    final_summary = [word.text for word in summary]

    summary = ''.join(final_summary)
    print("\nfinal summary:\n"+summary)

    # write text into txt
    with open(summarized_doc_path+"\\"+fn[0]+".txt", 'w') as f:
        for line in final_summary:
            f.write(line)
            f.write('\n')
    
if __name__ == '__main__':
    summerize("originalText.docx")